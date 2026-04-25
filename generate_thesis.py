#!/usr/bin/env python3
"""Generate thesis.docx from markdown draft, applying TYUT formatting rules."""

import re
import pathlib
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MD_PATH = pathlib.Path("/Users/stupidcat/Documents/T-Match/论文要求/论文初稿/简历岗位智能匹配与优化系统的设计与实现.md")
OUT_PATH = pathlib.Path("/Users/stupidcat/Documents/T-Match/thesis.docx")

CN_NORMAL = "宋体"
CN_HEADING = "黑体"
EN_FONT = "Times New Roman"
CODE_FONT = "Consolas"

# ─── Helpers ──────────────────────────────────────────────────────────────────

def set_cn_font(run, cn_name, en_name=EN_FONT, size=None):
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), cn_name)
    rFonts.set(qn('w:ascii'), en_name)
    rFonts.set(qn('w:hAnsi'), en_name)
    if size: run.font.size = size

def set_first_line_indent_chars(paragraph, chars=2):
    pPr = paragraph._element.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    ind.set(qn('w:firstLineChars'), str(chars * 100))

def set_page_margins(section):
    section.top_margin = Cm(3.3)
    section.bottom_margin = Cm(2.3)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.3)
    section.gutter = Cm(0)

def add_page_number_footer(section, fmt="decimal", start=None):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.font.size = Pt(9)
    set_cn_font(run, CN_NORMAL)
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar1)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run._element.append(instrText)
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._element.append(fldChar2)
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        pgNumType = OxmlElement('w:pgNumType')
        sectPr.append(pgNumType)
    pgNumType.set(qn('w:fmt'), fmt)
    if start: pgNumType.set(qn('w:start'), str(start))

def add_toc_field(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    set_cn_font(run, CN_NORMAL, size=Pt(12))
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar1)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    run._element.append(instrText)
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run._element.append(fldChar2)
    run2 = p.add_run("请右键此处更新目录")
    set_cn_font(run2, CN_NORMAL, size=Pt(12))
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run2._element.append(fldChar3)

def configure_styles(doc):
    normal = doc.styles['Normal']
    normal.font.size = Pt(12)
    set_cn_font(normal.font, CN_NORMAL)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    h1 = doc.styles['Heading 1']
    h1.font.size = Pt(15)
    h1.font.bold = True
    set_cn_font(h1.font, CN_HEADING)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    h1.paragraph_format.line_spacing = 1.5
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.keep_with_next = True

    h2 = doc.styles['Heading 2']
    h2.font.size = Pt(14)
    h2.font.bold = True
    set_cn_font(h2.font, CN_HEADING)
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    h2.paragraph_format.line_spacing = 1.5
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True

    h3 = doc.styles['Heading 3']
    h3.font.size = Pt(12)
    h3.font.bold = True
    set_cn_font(h3.font, CN_HEADING)
    h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    h3.paragraph_format.line_spacing = 1.5
    h3.paragraph_format.space_before = Pt(6)
    h3.paragraph_format.space_after = Pt(3)
    h3.paragraph_format.keep_with_next = True

# ─── Inline parsing ───────────────────────────────────────────────────────────

def parse_inline(paragraph, text):
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`[^`]+`)', text)
    for part in parts:
        if not part: continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            set_cn_font(run, CN_NORMAL, size=Pt(12))
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            set_cn_font(run, CN_NORMAL, size=Pt(12))
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            set_cn_font(run, CN_NORMAL, CODE_FONT, Pt(10.5))
        else:
            run = paragraph.add_run(part)
            set_cn_font(run, CN_NORMAL, size=Pt(12))

# ─── Markdown renderer ────────────────────────────────────────────────────────

HEADING1_CHAPTERS = {'摘  要', '摘要', 'Abstract', '目  录', '目录', '结论',
                     '参考文献', '致  谢', '致谢', '外文原文', '中文翻译'}

def is_chapter_heading(text):
    stripped = text.strip()
    if stripped in HEADING1_CHAPTERS: return True
    if re.match(r'^\d+\s', stripped): return True
    return False

def render_markdown(doc, md_text):
    lines = md_text.split('\n')
    state = 'normal'
    code_buffer = []
    table_buffer = []
    para_buffer = []
    first_h1_seen = False
    in_toc = False
    toc_ended = False

    for i, raw_line in enumerate(lines):
        line = raw_line.rstrip()

        # ── code block state ──
        if state == 'code':
            if line.startswith('```'):
                state = 'normal'
                add_code_block(doc, '\n'.join(code_buffer))
                code_buffer = []
                continue
            code_buffer.append(raw_line)
            continue

        # ── table state ──
        if state == 'table':
            if line.startswith('|'):
                row = [c.strip() for c in line.strip('|').split('|')]
                if not all(c in {'-', ':', ''} for c in row): table_buffer.append(row)
                continue
            else:
                state = 'normal'
                flush_table(doc, table_buffer)
                table_buffer = []

        # ── skip horizontal rules ──
        if re.match(r'^---+\s*$', line): continue

        # ── TOC handling ──
        if in_toc and not toc_ended:
            if line.startswith('## ') and not line.strip() in {'目  录', '目录'}:
                toc_ended = True
                in_toc = False
            else:
                if line.strip() and not line.startswith('#'): continue

        # ── heading detection ──
        if line.startswith('#'):
            level = len(re.match(r'^#+', line).group())
            heading_text = line.lstrip('#').strip()

            # flush paragraph buffer
            if para_buffer:
                add_normal_paragraph(doc, '\n'.join(para_buffer))
                para_buffer = []

            if level == 1 and not first_h1_seen:
                # thesis title
                first_h1_seen = True
                add_title(doc, heading_text)
                continue

            if level == 2:
                if heading_text in {'目  录', '目录'}:
                    in_toc = True
                    p = doc.add_heading(heading_text, level=1)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    add_toc_field(doc)
                    continue
                if is_chapter_heading(heading_text):
                    p = doc.add_heading(heading_text, level=1)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    set_first_line_indent_chars(p, 0)
                    continue
                else:
                    p = doc.add_heading(heading_text, level=2)
                    set_first_line_indent_chars(p, 0)
                    continue

            if level == 3:
                p = doc.add_heading(heading_text, level=2)
                set_first_line_indent_chars(p, 0)
                continue

            if level >= 4:
                p = doc.add_heading(heading_text, level=3)
                set_first_line_indent_chars(p, 0)
                continue

        # ── fenced code block start ──
        if line.startswith('```'):
            if para_buffer:
                add_normal_paragraph(doc, '\n'.join(para_buffer))
                para_buffer = []
            state = 'code'
            continue

        # ── GFM table start ──
        if line.startswith('|') and state == 'normal':
            if para_buffer:
                add_normal_paragraph(doc, '\n'.join(para_buffer))
                para_buffer = []
            row = [c.strip() for c in line.strip('|').split('|')]
            if not all(c in {'-', ':', ''} for c in row): table_buffer.append(row)
            state = 'table'
            continue

        # ── blank line → flush paragraph ──
        if not line.strip():
            if para_buffer:
                add_normal_paragraph(doc, '\n'.join(para_buffer))
                para_buffer = []
            continue

        # ── special: keywords line ──
        if line.strip().startswith('**关键词') or line.strip().startswith('**Key words'):
            if para_buffer:
                add_normal_paragraph(doc, '\n'.join(para_buffer))
                para_buffer = []
            add_keywords_paragraph(doc, line.strip())
            continue

        # ── special: table caption like **表4-1 xxx** ──
        if re.match(r'^\*?表\d+-\d+\s', line.strip()) or re.match(r'^\*?图\d+-\d+\s', line.strip()):
            if para_buffer:
                add_normal_paragraph(doc, '\n'.join(para_buffer))
                para_buffer = []
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            text = re.sub(r'^\*+', '', line.strip()).rstrip('*')
            run = p.add_run(text)
            set_cn_font(run, CN_NORMAL, size=Pt(12))
            continue

        # ── accumulate paragraph ──
        para_buffer.append(line)

    # flush remaining
    if para_buffer: add_normal_paragraph(doc, '\n'.join(para_buffer))
    if state == 'code': add_code_block(doc, '\n'.join(code_buffer))
    if state == 'table' and table_buffer: flush_table(doc, table_buffer)

# ─── Element builders ─────────────────────────────────────────────────────────

def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)
    p.paragraph_format.space_after = Pt(24)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.size = Pt(22)
    run.bold = True
    set_cn_font(run, CN_HEADING)
    set_first_line_indent_chars(p, 0)

def add_normal_paragraph(doc, text):
    text = text.strip()
    if not text: return
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    set_first_line_indent_chars(p, 2)
    parse_inline(p, text)

def add_keywords_paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    set_first_line_indent_chars(p, 0)
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if not part: continue
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
            set_cn_font(run, CN_NORMAL, size=Pt(12))
        else:
            run = p.add_run(part)
            set_cn_font(run, CN_NORMAL, size=Pt(12))

def add_code_block(doc, code):
    lines = code.split('\n')
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.5)
        set_first_line_indent_chars(p, 0)
        run = p.add_run(line)
        run.font.size = Pt(10.5)
        set_cn_font(run, CN_NORMAL, CODE_FONT, Pt(10.5))
        run.italic = True

def flush_table(doc, rows):
    if len(rows) < 2: return
    ncols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            if c_idx >= ncols: break
            cell = table.cell(r_idx, c_idx)
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(cell_text)
            run.font.size = Pt(10.5)
            set_cn_font(run, CN_NORMAL, size=Pt(10.5))
            if r_idx == 0: run.bold = True

# ─── Main ─────────────────────────────────────────────────────────────────────

def main():
    md = MD_PATH.read_text(encoding="utf-8")
    doc = Document()
    configure_styles(doc)

    # Initial section setup
    section = doc.sections[0]
    set_page_margins(section)
    add_page_number_footer(section, fmt="decimal", start=1)

    render_markdown(doc, md)
    doc.save(str(OUT_PATH))
    print(f"Generated {OUT_PATH}")

if __name__ == "__main__":
    main()